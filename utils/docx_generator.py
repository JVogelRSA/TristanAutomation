import io
import re
import unicodedata
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from htmldocx import HtmlToDocx


# ── Colour palette ─────────────────────────────────────────────
NAVY        = (0x1E, 0x3A, 0x5F)
DARK_GRAY   = (0x44, 0x44, 0x44)
MID_GRAY    = (0x99, 0x99, 0x99)
RED_BG      = (0xFF, 0xE0, 0xDC)
ORANGE_BG   = (0xFF, 0xF3, 0xCC)
GREEN_BG    = (0xD5, 0xF5, 0xE3)
BLUE_BG     = (0xEB, 0xF5, 0xFB)
WHITE       = (0xFF, 0xFF, 0xFF)
RED_TEXT    = (0xA9, 0x33, 0x26)
ORANGE_TEXT = (0xB7, 0x60, 0x0E)
GREEN_TEXT  = (0x1A, 0x7A, 0x42)

NAVY_RGB  = RGBColor(*NAVY)
WHITE_RGB = RGBColor(*WHITE)
GRAY_RGB  = RGBColor(*MID_GRAY)


# ── Emoji / artefact cleaning ──────────────────────────────────
_EMOJI_MAP = {
    # Decorative section icons → stripped
    '\U0001F31F': '',   # 🌟
    '\U0001F4C5': '',   # 📅
    '\U000023F3': '',   # ⏳
    '\U0001F4CB': '',   # 📋
    '\U0001F4B0': '',   # 💰
    '\U0001F4CA': '',   # 📊
    '\U0001F50D': '',   # 🔍
    '\U0001F6A8': '',   # 🚨
    '\U0001F4A1': '',   # 💡
    '\U0001F9FE': '',   # 🧾
    # Status circles → text labels
    '\U0001F534': '[ALERT]',   # 🔴
    '\U0001F7E2': '[OK]',      # 🟢
    '\U0001F7E1': '[WARN]',    # 🟡
}


def _strip_emoji(text: str) -> str:
    for ch, repl in _EMOJI_MAP.items():
        text = text.replace(ch, repl)
    # Drop any remaining non-BMP / symbol characters that DOCX can't render
    return ''.join(
        c for c in text
        if unicodedata.category(c) not in ('So', 'Cs')
    )


def _clean_html(html: str) -> str:
    """Strip code fences, stray markdown bold, and unsafe emoji."""
    html = html.strip()
    if html.startswith('```html'):
        html = html[7:]
    elif html.startswith('```'):
        html = html[3:]
    if html.endswith('```'):
        html = html[:-3]
    html = html.strip()
    # Convert **markdown** bold → <b>HTML bold</b>
    html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', html, flags=re.DOTALL)
    html = _strip_emoji(html)
    return html


# ── python-docx helpers ────────────────────────────────────────
def _set_cell_shading(cell, rgb_tuple):
    """Apply a solid background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shading element first
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb_tuple))
    tcPr.append(shd)


def _set_cell_text_color(cell, rgb_tuple):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(*rgb_tuple)


def _get_runway_weeks(text: str):
    """Return numeric runway if the text contains a '# weeks' pattern."""
    m = re.search(r'(\d+\.?\d*)\s*week', text, re.IGNORECASE)
    if m:
        return float(m.group(1))
    return None


def _classify_row(cells):
    """
    Decide the highlight colour for a data row.
    Returns (background_rgb, text_rgb) or (None, None) for zebra-only styling.
    """
    combined = ' '.join(c.text for c in cells)
    lower    = combined.lower()

    # Runway-based (inventory)
    runway = _get_runway_weeks(combined)
    if runway is not None:
        if runway < 4:
            return RED_BG, RED_TEXT
        if runway < 8:
            return ORANGE_BG, ORANGE_TEXT
        return GREEN_BG, GREEN_TEXT

    # Keyword-based critical flags
    if any(k in lower for k in ('critical', 'out of stock', 'stockout', '[alert]')):
        return RED_BG, RED_TEXT
    if any(k in lower for k in ('warning', 'reorder', 'low stock', '[warn]')):
        return ORANGE_BG, ORANGE_TEXT

    # Large spend anomalies (spend report)
    for amt_str in re.findall(r'\$[\d,]+\.?\d*', combined):
        try:
            amt = float(amt_str.replace('$', '').replace(',', ''))
            if amt >= 5000:
                return ORANGE_BG, ORANGE_TEXT
        except ValueError:
            pass

    return None, None


def _style_all_tables(doc):
    """
    Post-process every table in the document:
      • Row 0  → navy header with white bold text
      • Rows 1+ → colour-coded by content, then zebra-striped
    """
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if i == 0:
                # Header
                for cell in row.cells:
                    _set_cell_shading(cell, NAVY)
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = WHITE_RGB
                            run.font.size = Pt(9)
            else:
                bg, fg = _classify_row(row.cells)
                for cell in row.cells:
                    _set_cell_shading(cell, bg if bg else (BLUE_BG if i % 2 == 0 else WHITE))
                    if fg:
                        _set_cell_text_color(cell, fg)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)


def _style_headings_and_body(doc):
    """
    Apply navy colour to heading paragraphs and cap body font to 10pt.
    """
    for para in doc.paragraphs:
        style = para.style.name.lower()
        if 'heading' in style:
            for run in para.runs:
                run.font.color.rgb = NAVY_RGB
                run.bold = True
                run.font.size = Pt(12 if '1' in style else 10.5)
        else:
            # Flag CRITICAL / ALERT paragraphs in red
            if any(k in para.text.upper() for k in ('CRITICAL ALERT', '[ALERT]', 'IMMEDIATE REORDER')):
                for run in para.runs:
                    run.font.color.rgb = RGBColor(*RED_TEXT)
                    run.bold = True
            for run in para.runs:
                if run.font.size is None or run.font.size > Pt(11):
                    run.font.size = Pt(10)


# ── Public API ─────────────────────────────────────────────────
def html_to_docx(html_content: str, title: str,
                 date_str: str = None,
                 chart_images: list = None) -> bytes:
    """
    Convert LLM-generated HTML into a polished, colour-coded DOCX report.

    Parameters
    ----------
    html_content : str
        Raw HTML from the LLM (may contain emoji, markdown artefacts, code fences).
    title : str
        Report title shown in the header block.
    date_str : str, optional
        ISO date string shown in the subtitle.
    chart_images : list[bytes], optional
        PNG image bytes to embed below the header, before the HTML body.

    Returns
    -------
    bytes
        DOCX file ready for email attachment.
    """
    if not date_str:
        date_str = datetime.now().strftime('%Y-%m-%d')

    html_content = _clean_html(html_content)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Title block ───────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title.upper())
    tr.bold = True
    tr.font.size = Pt(17)
    tr.font.color.rgb = NAVY_RGB
    tp.paragraph_format.space_after = Pt(2)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(
        f"Week of {date_str}   \u00b7   Generated {datetime.now().strftime('%B %d, %Y')}"
    )
    dr.font.size = Pt(9)
    dr.font.color.rgb = GRAY_RGB
    dp.paragraph_format.space_after = Pt(6)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run('\u2500' * 90)
    rr.font.size = Pt(7)
    rr.font.color.rgb = NAVY_RGB
    rule.paragraph_format.space_after = Pt(10)

    # ── Charts ────────────────────────────────────────────────
    if chart_images:
        for img_bytes in chart_images:
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    # ── HTML body ─────────────────────────────────────────────
    parser = HtmlToDocx()
    parser.table_style = 'Table Grid'
    parser.add_html_to_document(html_content, doc)

    # ── Post-process ──────────────────────────────────────────
    _style_headings_and_body(doc)
    _style_all_tables(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
